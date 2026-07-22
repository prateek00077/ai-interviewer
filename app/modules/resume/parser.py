"""PDF/DOCX -> structured resume fields.

Deliberately not an LLM call. Extraction here is mechanical -- pull the text,
find the section headings, pick contact details out with regexes -- and a model
would add a second or two of latency, a per-upload cost, and a way for the
document to talk back. A resume is untrusted input authored by the person being
evaluated; text arriving from it must never become an instruction.

The parse is best-effort by design. Resumes have no schema, and a document whose
sections cannot be identified still yields usable text: ``sections`` falls back
to one "body" span rather than failing. Only a document with no extractable text
at all is a parse failure.
"""

from __future__ import annotations

import io
import re
import unicodedata
from dataclasses import dataclass, field

import structlog

log = structlog.get_logger(__name__)

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Enough text to be a CV rather than a blank scan. A scanned image PDF extracts
# to roughly nothing, and OCR is out of scope, so this is where that fails.
MIN_EXTRACTED_CHARS = 120


class ResumeParseError(Exception):
    """Unparseable document. The message is operator-facing, not candidate-facing."""


# Canonical section name -> the headings that introduce it. Matched against a
# whole line, case-insensitively, so "SKILLS" and "Technical Skills:" both land
# in "skills".
SECTION_PATTERNS: dict[str, tuple[str, ...]] = {
    "summary": ("summary", "objective", "profile", "about"),
    "experience": (
        "experience",
        "work experience",
        "employment",
        "professional experience",
        "work history",
    ),
    "education": ("education", "academic", "qualifications"),
    "skills": ("skills", "technical skills", "technologies", "competencies"),
    "projects": ("projects", "personal projects", "selected projects"),
    "certifications": ("certifications", "certificates", "licenses"),
    "awards": ("awards", "honors", "achievements"),
    "publications": ("publications", "papers"),
}

# A heading is a short line, so a sentence merely containing the word
# "experience" is not mistaken for the start of a section.
MAX_HEADING_CHARS = 60

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
# Deliberately loose: international formats vary too much for a strict pattern,
# and a false positive here costs nothing -- the value is informational.
PHONE_RE = re.compile(r"(?:\+\d{1,3}[\s-]?)?(?:\(?\d{2,4}\)?[\s-]?){2,4}\d{2,4}")
URL_RE = re.compile(r"https?://[^\s<>\"]+|(?:www\.|linkedin\.com/|github\.com/)[^\s<>\"]+")


@dataclass(slots=True)
class ParsedResume:
    text: str
    sections: dict[str, str] = field(default_factory=dict)
    emails: list[str] = field(default_factory=list)
    phones: list[str] = field(default_factory=list)
    links: list[str] = field(default_factory=list)

    def as_dict(self) -> dict:
        """The shape stored in ``Resume.parsed``.

        Full text is excluded: it is already the source of the chunks, and
        duplicating it into JSONB doubles the row size for no reader.
        """
        return {
            "sections": self.sections,
            "emails": self.emails,
            "phones": self.phones,
            "links": self.links,
            "char_count": len(self.text),
        }


# --- Extraction -------------------------------------------------------------


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    # PyPdfError is the base of every pypdf failure; PdfReadError alone misses
    # the stream and dependency errors a malformed file can raise.
    from pypdf.errors import PyPdfError

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # An empty password unlocks most "protected" PDFs people export.
            try:
                reader.decrypt("")
            except Exception as exc:  # noqa: BLE001 - any failure means unreadable
                raise ResumeParseError("PDF is password protected") from exc
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except PyPdfError as exc:
        raise ResumeParseError(f"Malformed PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(io.BytesIO(data))
    except (PackageNotFoundError, KeyError, ValueError) as exc:
        raise ResumeParseError(f"Malformed DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    # Tables are common in resume templates and are invisible to `paragraphs`.
    # Skipping them silently drops whole employment histories.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


EXTRACTORS = {PDF_CONTENT_TYPE: _extract_pdf, DOCX_CONTENT_TYPE: _extract_docx}


def normalize(raw: str) -> str:
    """Collapse the noise a PDF text layer produces.

    NFKC folds ligatures and full-width characters that would otherwise split a
    word into two tokens. Bullets become newlines so each becomes its own line.
    """
    text = unicodedata.normalize("NFKC", raw)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[•●▪‣⁃]", "\n", text)
    text = re.sub(r"[ \t ]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(line.strip() for line in text.split("\n")).strip()


# --- Sectioning -------------------------------------------------------------


def _heading_for(line: str) -> str | None:
    """The canonical section a line introduces, if it is a heading at all."""
    stripped = line.strip().strip(":").strip()
    if not stripped or len(stripped) > MAX_HEADING_CHARS:
        return None
    lowered = stripped.lower()
    for canonical, aliases in SECTION_PATTERNS.items():
        if lowered in aliases:
            return canonical
    return None


def split_sections(text: str) -> dict[str, str]:
    """Group lines under the last heading seen.

    Anything before the first heading is "header" -- on a real CV that is the
    name and contact block, which is worth keeping separate from the summary.
    """
    sections: dict[str, list[str]] = {}
    current = "header"
    for line in text.split("\n"):
        heading = _heading_for(line)
        if heading is not None:
            current = heading
            sections.setdefault(current, [])
            continue
        if line:
            sections.setdefault(current, []).append(line)

    collapsed = {name: "\n".join(lines).strip() for name, lines in sections.items()}
    return {name: body for name, body in collapsed.items() if body}


# --- Contact details --------------------------------------------------------


def _dedupe(values: list[str]) -> list[str]:
    """Order-preserving, case-insensitive."""
    seen: set[str] = set()
    out: list[str] = []
    for value in values:
        key = value.lower()
        if key not in seen:
            seen.add(key)
            out.append(value)
    return out


def extract_contacts(text: str) -> tuple[list[str], list[str], list[str]]:
    emails = _dedupe(EMAIL_RE.findall(text))
    links = _dedupe(URL_RE.findall(text))

    # Run the phone pattern over text with emails and URLs removed: both are full
    # of digit runs that otherwise match as phone numbers.
    scrubbed = URL_RE.sub(" ", EMAIL_RE.sub(" ", text))
    phones = [
        match.strip()
        for match in PHONE_RE.findall(scrubbed)
        # At least 8 digits, so dates and postcodes do not qualify.
        if len(re.sub(r"\D", "", match)) >= 8
    ]
    return emails, _dedupe(phones), links


# --- Entry point ------------------------------------------------------------


def parse(data: bytes, content_type: str) -> ParsedResume:
    """Bytes to structured fields. Raises ResumeParseError on unusable input."""
    extractor = EXTRACTORS.get(content_type)
    if extractor is None:
        raise ResumeParseError(f"Unsupported content type: {content_type}")

    text = normalize(extractor(data))
    if len(text) < MIN_EXTRACTED_CHARS:
        # Almost always a scanned image with no text layer. OCR is out of scope,
        # and silently storing an empty resume would poison retrieval later.
        raise ResumeParseError(
            f"Extracted only {len(text)} characters; the document may be a scan"
        )

    emails, phones, links = extract_contacts(text)
    sections = split_sections(text) or {"body": text}

    log.info(
        "resume_parsed",
        content_type=content_type,
        chars=len(text),
        sections=sorted(sections),
    )
    return ParsedResume(
        text=text, sections=sections, emails=emails, phones=phones, links=links
    )
